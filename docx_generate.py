from user_data import Sent_applications, Institution
import os
import docx
import re
from template_parsing import teplate_parsing

def rename_docx_file(file_path, title_olympiad, name, surname, patronymic, institution_director, institutions_name, classes):
    new_file_name = f"{title_olympiad} {name} {surname} {patronymic}.docx"
    new_file_path = os.path.join(os.path.dirname(file_path), new_file_name)

    institution_director_short = institution_director.split()
    classe = classes.split()

    dative_surname, respect = teplate_parsing(institution_director_short, classe)



    director_EI =f'{institution_director_short[1][0]}.{institution_director_short[2][0]}. {dative_surname}'
    dictionary = {"{appeal_to_director}": f'{respect} {institution_director_short[1]} {institution_director_short[2]}',
                  "{class_number}": classe[0],
                  "{{EI}}": institutions_name,
                  "{{director_EI}}": director_EI,
                  "{FIO_participant}":  f"{surname} {name} {patronymic}"}


    doc = docx.Document(file_path)

    style = doc.styles['Normal']
    font = style.font
    font.size = docx.shared.Pt(14)
    font.name = 'Times New Roman'

    #Сори за это, ты говорил ещё давно ,чтобы не было два цикла, а тут их ещё больше
    #Но там в шаблоне таблица, как по другому я хз

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for pattern, replacement in dictionary.items():
                        if re.search(pattern, paragraph.text):
                            paragraph.text = re.sub(pattern, replacement, paragraph.text)

    #Тут тож хз как сделать 1 цикл
    for paragraph in doc.paragraphs:
        for pattern, replacement in dictionary.items():
            if re.search(pattern, paragraph.text):
                paragraph.text = re.sub(pattern, replacement, paragraph.text)


    doc.save(new_file_path)

    return new_file_path







